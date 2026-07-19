"""
多格式文档解析服务
支持 PDF、Word、HTML、Markdown、TXT 等格式的文本提取
"""
import os
import re
import json
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from abc import ABC, abstractmethod

# PDF 解析
try:
    import pdfplumber
    PDF_PARSER = "pdfplumber"
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_PARSER = "PyPDF2"
    except ImportError:
        PDF_PARSER = None

# Word 解析
try:
    from docx import Document as DocxDocument
    DOCX_PARSER = "python-docx"
except ImportError:
    DOCX_PARSER = None

# HTML 解析
try:
    from bs4 import BeautifulSoup
    HTML_PARSER = "beautifulsoup"
except ImportError:
    HTML_PARSER = None

# Markdown 处理
try:
    import markdown
    MARKDOWN_PARSER = "markdown"
except ImportError:
    MARKDOWN_PARSER = None


class DocumentParser(ABC):
    """文档解析基类"""

    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析文档，返回包含文本内容和元数据的字典"""
        pass

    @abstractmethod
    def extract_text(self, content: Any) -> str:
        """从文档内容中提取纯文本"""
        pass


class PDFParser(DocumentParser):
    """PDF文档解析器"""

    def __init__(self):
        self.parser_name = PDF_PARSER

    def _extract_with_font_sizes(self, file_path: str) -> tuple:
        """使用PyMuPDF提取文本及其字体大小，按行聚合"""
        import fitz
        lines_data = []  # [(text, dominant_font_size, is_centered), ...]
        metadata = {"pages": 0, "title": "", "author": ""}

        doc = fitz.open(file_path)
        metadata["pages"] = len(doc)

        for page in doc:
            text_dict = page.get_text("dict")
            page_width = page.rect.width

            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:  # 只处理文本块
                    continue

                bbox = block.get("bbox", [])
                block_width = (bbox[2] - bbox[0]) if len(bbox) == 4 else 0
                is_centered = block_width < page_width * 0.6 if page_width > 0 else False

                for line in block.get("lines", []):
                    line_texts = []
                    font_sizes = []

                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        font_size = span.get("size", 0)
                        if text.strip():
                            line_texts.append(text)
                            font_sizes.append(font_size)

                    if line_texts:
                        # 取这行最常见的字体大小作为 dominant size
                        text = "".join(line_texts)
                        if font_sizes:
                            # 使用中位数而不是平均数，更抗噪声
                            sorted_sizes = sorted(font_sizes)
                            mid = len(sorted_sizes) // 2
                            dominant_size = sorted_sizes[mid] if len(sorted_sizes) % 2 == 1 else sorted_sizes[mid - 1]
                        else:
                            dominant_size = 0

                        lines_data.append({
                            "text": text.strip(),
                            "font_size": dominant_size,
                            "is_centered": is_centered,
                            "page": page.number + 1
                        })

        doc.close()
        return lines_data, metadata

    def _detect_headings_by_font_size(self, lines_data: list) -> list:
        """根据字体大小检测标题，返回带标题标记的文本"""
        if not lines_data:
            return []

        # 收集所有字体大小
        font_sizes = [d["font_size"] for d in lines_data if d["font_size"] > 0]
        if not font_sizes:
            return [d["text"] for d in lines_data]

        # 统计字体大小分布
        size_counts = {}
        for fs in font_sizes:
            # 四舍五入到整数，避免微小的浮点差异
            rounded = round(fs)
            size_counts[rounded] = size_counts.get(rounded, 0) + 1

        # 找出最常见的字体大小（正文大小）
        max_count = max(size_counts.values())
        body_sizes = [size for size, count in size_counts.items() if count >= max_count * 0.3]
        body_size = max(body_sizes) if body_sizes else max(size_counts.keys())

        # 找出比正文明显大的字体大小（可能是标题）
        heading_sizes = sorted([size for size in size_counts.keys() if size > body_size], reverse=True)

        result = []
        for line in lines_data:
            text = line["text"]

            # 跳过纯页码
            if re.match(r'^\d+$', text):
                continue

            font_size = round(line["font_size"])

            # 判断是否是标题
            is_heading = font_size > body_size
            is_centered_heading = line["is_centered"] and len(text) < 80 and len(text) > 2

            if is_heading or is_centered_heading:
                # 根据字体大小确定标题级别
                if heading_sizes and font_size >= heading_sizes[0]:
                    level = 1
                elif len(heading_sizes) > 1 and font_size >= heading_sizes[1]:
                    level = 2
                else:
                    level = 2  # 默认二级标题

                result.append(f"{'#' * level} {text}")
            else:
                result.append(text)

        return result

    def _normalize_text(self, text: str) -> str:
        """规范化PDF提取的文本，清理多余空行"""
        lines = text.split('\n')
        result_lines = []
        prev_empty = False

        for line in lines:
            line = line.strip()
            is_empty = not line

            # 跳过连续空行
            if is_empty:
                if not prev_empty:
                    result_lines.append('')
                prev_empty = True
                continue

            prev_empty = False
            result_lines.append(line)

        # 清理首尾空行
        while result_lines and not result_lines[0]:
            result_lines.pop(0)
        while result_lines and not result_lines[-1]:
            result_lines.pop()

        return '\n'.join(result_lines)

    def _extract_title_from_content(self, content: str) -> str:
        """从内容中提取可能的标题（已禁用）"""
        return ""

    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if PDF_PARSER is None and PDF_PARSER != "pdfplumber":
            raise ImportError("请安装 pdfplumber: pip install pdfplumber")

        metadata = {"pages": 0, "title": "", "author": ""}
        parse_error = None
        content = ""

        # 优先使用 pdfplumber（稳定可靠）
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                if pdf.metadata:
                    metadata["title"] = pdf.metadata.get("Title", "")
                    metadata["author"] = pdf.metadata.get("Author", "")

                for i, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text)
                    except Exception as e:
                        print(f"提取PDF第{i+1}页文本失败: {e}")
                        continue

            if text_parts:
                content = self._normalize_text('\n\n'.join(text_parts))
            else:
                raise Exception("pdfplumber未能提取到文本")

        except Exception as e:
            parse_error = str(e)
            print(f"pdfplumber解析PDF失败: {e}")

            # 尝试 PyPDF2 作为后备
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                metadata["pages"] = len(reader.pages)
                if reader.metadata:
                    metadata["title"] = metadata["title"] or reader.metadata.get("/Title", "")
                    metadata["author"] = reader.metadata.get("/Author", "")

                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)

                if text_parts:
                    content = self._normalize_text('\n\n'.join(text_parts))
                else:
                    raise Exception("PyPDF2未能提取到文本")
            except Exception as e2:
                print(f"PyPDF2解析PDF失败: {e2}")

        # 如果仍然没有提取到文本
        if not content.strip():
            if parse_error:
                raise Exception(f"PDF解析失败: {parse_error}，且无后备解析器可用。请尝试将PDF转换为文本格式后重新上传。")
            content = "[此PDF为扫描版或图片型PDF，无法提取文本内容。请尝试使用文字型PDF重新上传。]"

        return {
            "content": content,
            "metadata": metadata,
            "source": os.path.basename(file_path),
            "format": "pdf"
        }

    def extract_text(self, content: Any) -> str:
        """提取纯文本"""
        return content


class DocxParser(DocumentParser):
    """Word文档解析器 - 输出Markdown格式"""

    def __init__(self):
        self.parser_name = DOCX_PARSER

    def _get_paragraph_text_with_format(self, para) -> str:
        """获取段落文本，保留行内格式（加粗、斜体等）"""
        from docx.oxml.ns import qn

        parts = []
        for child in para._p:
            # 处理加粗
            if child.tag.endswith('}r'):
                # 检查是否有加粗标记
                rPr = child.find(qn('w:rPr'))
                is_bold = rPr is not None and rPr.find(qn('w:b')) is not None
                is_italic = rPr is not None and rPr.find(qn('w:i')) is not None

                text = ''.join(t.text for t in child.findall(qn('w:t')) if t.text)

                if text:
                    if is_bold and is_italic:
                        parts.append(f"***{text}***")
                    elif is_bold:
                        parts.append(f"**{text}**")
                    elif is_italic:
                        parts.append(f"*{text}*")
                    else:
                        parts.append(text)

        return ''.join(parts)

    def _paragraph_to_markdown(self, para) -> str:
        """将段落转换为Markdown格式"""
        style_name = para.style.name.lower() if para.style else ""

        # 处理标题样式
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
                level = min(max(level, 1), 6)  # 限制在1-6级
            except ValueError:
                level = 1
            text = para.text.strip()
            return f"{'#' * level} {text}\n" if text else ""

        # 处理列表项
        if style_name in ["list bullet", "list bullet 2", "list bullet 3"]:
            text = para.text.strip()
            return f"- {text}\n" if text else ""

        if style_name in ["list number", "list number 2", "list number 3"]:
            # 简单处理，输出为有序列表
            text = para.text.strip()
            return f"1. {text}\n" if text else ""

        # 普通段落
        text = self._get_paragraph_text_with_format(para)
        return f"{text}\n\n" if text.strip() else ""

    def _table_to_markdown(self, table) -> str:
        """将表格转换为Markdown格式"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip().replace('\n', ' '))
            rows.append(cells)

        if not rows:
            return ""

        # 确定列数
        col_count = len(rows[0]) if rows else 0
        if col_count == 0:
            return ""

        md_rows = []
        for idx, row in enumerate(rows):
            md_rows.append("| " + " | ".join(row) + " |")
            if idx == 0:
                md_rows.append("| " + " | ".join(["---"] * col_count) + " |")

        return "\n".join(md_rows) + "\n\n"

    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析Word文档，输出Markdown格式"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if DOCX_PARSER is None:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = DocxDocument(file_path)

        # 提取标题 - 优先查找一级标题
        title = ""
        heading_priority = ["heading 1", "标题 1", "heading1", "标题1", "title"]

        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ""

            # 一级标题优先
            if style_name in heading_priority:
                title = para.text.strip()
                break

        # 如果没找到一级标题，找任意标题
        if not title:
            for para in doc.paragraphs:
                style_name = para.style.name.lower() if para.style else ""
                if style_name.startswith("heading") or style_name.startswith("标题"):
                    title = para.text.strip()
                    if title:
                        break

        # 如果还是没有，使用文件名
        if not title:
            title = os.path.basename(file_path).rsplit('.', 1)[0]

        # 转换为Markdown格式
        content_parts = []
        in_table = False

        for element in doc.element.body:
            # 判断元素类型
            if element.tag.endswith('p'):  # 段落
                # 找到对应的paragraph对象
                para = None
                for p in doc.paragraphs:
                    if p._p is element:
                        para = p
                        break

                if para is None:
                    continue

                style_name = para.style.name.lower() if para.style else ""

                # 如果是表格内的段落，跳过
                if in_table:
                    in_table = False
                    continue

                md = self._paragraph_to_markdown(para)
                if md.strip():
                    content_parts.append(md)

            elif element.tag.endswith('tbl'):  # 表格
                # 找到对应的table对象
                table_obj = None
                for t in doc.tables:
                    if t._tbl is element:
                        table_obj = t
                        break

                if table_obj:
                    table_md = self._table_to_markdown(table_obj)
                    if table_md.strip():
                        content_parts.append(table_md)

        content = "".join(content_parts)

        # 清理多余空行
        content = re.sub(r'\n{3,}', '\n\n', content)

        return {
            "content": content.strip(),
            "metadata": {
                "title": title,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "author": ""
            },
            "source": os.path.basename(file_path),
            "format": "docx"
        }

    def extract_text(self, content: Any) -> str:
        """提取纯文本"""
        return content


class HTMLParser(DocumentParser):
    """HTML文档解析器 - 输出Markdown格式"""

    def __init__(self):
        self.parser_name = HTML_PARSER

    def _get_inner_html(self, element):
        """获取元素的内部HTML"""
        return ''.join(str(child) for child in element.children)

    def _convert_inline_tags(self, element):
        """转换行内格式标签为Markdown"""
        text = element.decode_contents() if hasattr(element, 'decode_contents') else str(element)
        # 转换 <strong> 和 <b> 为 **text**
        text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text, flags=re.DOTALL)
        text = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', text, flags=re.DOTALL)
        # 转换 <em> 和 <i> 为 *text*
        text = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', text, flags=re.DOTALL)
        text = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', text, flags=re.DOTALL)
        # 转换 <code> 为 `text`
        text = re.sub(r'<code[^>]*>(.*?)</code>', r'`\1`', text, flags=re.DOTALL)
        # 转换 <a> 为 [text](url)
        text = re.sub(r'<a[^>]*href=["\']([^"\']*)["\'][^>]*>(.*?)</a>', r'[\2](\1)', text, flags=re.DOTALL)
        # 移除剩余的HTML标签
        text = re.sub(r'<[^>]+>', '', text)
        # 解码HTML实体
        text = text.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        text = text.replace('&#39;', "'").replace('&quot;', '"')
        return text.strip()

    def _parse_element(self, element):
        """递归解析元素，返回Markdown格式"""
        parts = []

        for child in element.children:
            if isinstance(child, str):
                text = child.strip()
                if text:
                    parts.append(text)
                continue

            tag_name = child.name.lower() if hasattr(child, 'name') else ''

            if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag_name[1])
                text = child.get_text(strip=True)
                if text:
                    parts.append(f"{'#' * level} {text}")
                    parts.append("")  # 空行

            elif tag_name == 'p':
                text = self._convert_inline_tags(child)
                if text:
                    parts.append(text)
                    parts.append("")  # 空行

            elif tag_name == 'ul':
                for li in child.find_all("li", recursive=False):
                    text = li.get_text(strip=True)
                    if text:
                        parts.append(f"- {text}")
                parts.append("")  # 空行

            elif tag_name == 'ol':
                for idx, li in enumerate(child.find_all("li", recursive=False), 1):
                    text = li.get_text(strip=True)
                    if text:
                        parts.append(f"{idx}. {text}")
                parts.append("")  # 空行

            elif tag_name == 'table':
                table_md = self._parse_table(child)
                if table_md:
                    parts.append(table_md)
                    parts.append("")  # 空行

            elif tag_name == 'pre':
                code = child.get_text(strip=True)
                if code:
                    parts.append(f"```\n{code}\n```")
                    parts.append("")  # 空行

            elif tag_name == 'blockquote':
                text = child.get_text(strip=True)
                if text:
                    for line in text.split('\n'):
                        parts.append(f"> {line}")
                    parts.append("")  # 空行

            elif tag_name in ['div', 'article', 'section', 'main', 'aside']:
                # 递归处理容器元素
                inner = self._parse_element(child)
                if inner:
                    parts.append(inner)

            elif tag_name in ['br', 'hr']:
                parts.append("")  # 换行

            elif tag_name in ['script', 'style', 'nav', 'header', 'footer', 'noscript']:
                continue  # 跳过

            elif tag_name == 'img':
                alt = child.get('alt', '')
                src = child.get('src', '')
                if src:
                    parts.append(f"![{alt}]({src})")

        return "\n".join(parts)

    def _parse_table(self, table):
        """解析表格为Markdown格式"""
        rows = []
        for tr in table.find_all("tr"):
            cells = []
            for td in tr.find_all(["td", "th"]):
                text = td.get_text(strip=True)
                cells.append(text)
            if cells:
                rows.append(cells)

        if not rows:
            return ""

        # 确定列数
        col_count = max(len(row) for row in rows) if rows else 0
        if col_count == 0:
            return ""

        # 生成Markdown表格
        md_rows = []
        for idx, row in enumerate(rows):
            md_rows.append("| " + " | ".join(row) + " |")
            if idx == 0:
                # 表头后的分隔行
                md_rows.append("| " + " | ".join(["---"] * col_count) + " |")

        return "\n".join(md_rows)

    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析HTML文件，输出Markdown格式"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if HTML_PARSER is None:
            raise ImportError("请安装 beautifulsoup4: pip install beautifulsoup4")

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # 移除不需要的元素
        for tag in soup(["script", "style", "nav", "header", "footer", "noscript"]):
            tag.decompose()

        # 提取标题
        title_tag = soup.find("title")
        title = title_tag.text.strip() if title_tag else ""

        # 提取 h1-h6 标题
        headings = []
        for i in range(1, 7):
            for h in soup.find_all(f"h{i}"):
                if h.text.strip():
                    headings.append({"level": i, "text": h.text.strip()})

        # 如果没有找到标题，尝试从第一个h1或body获取
        if not title and headings:
            title = headings[0]["text"]
        elif not title:
            first_h = soup.find(["h1", "h2", "h3"])
            if first_h:
                title = first_h.text.strip()

        # 解析body或整个文档
        body = soup.find("body")
        parse_root = body if body else soup

        # 解析为Markdown
        content = self._parse_element(parse_root)

        # 清理多余空行
        content = re.sub(r'\n{3,}', '\n\n', content)
        content = content.strip()

        # 统计信息
        paragraphs_count = len([p for p in content.split('\n\n') if p and not p.startswith('#') and not p.startswith('-') and not p.startswith('|')])

        return {
            "content": content,
            "metadata": {
                "title": title,
                "headings": headings,
                "paragraphs_count": paragraphs_count
            },
            "source": os.path.basename(file_path),
            "format": "html"
        }

    def extract_text(self, content: Any) -> str:
        """提取纯文本"""
        return content


class MarkdownParser(DocumentParser):
    """Markdown文档解析器"""

    def __init__(self):
        self.parser_name = MARKDOWN_PARSER

    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析Markdown文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        with open(file_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        # 提取YAML front matter
        front_matter = {}
        if md_content.startswith("---"):
            parts = md_content.split("---", 2)
            if len(parts) >= 3:
                try:
                    import yaml
                    front_matter = yaml.safe_load(parts[1]) or {}
                    md_content = parts[2]
                except ImportError:
                    pass

        # 提取标题（保留原始内容，不移除格式）
        title = ""
        lines = md_content.split("\n")
        for line in lines:
            if line.startswith("# "):
                title = line[2:].strip()
                break

        # 清理多余空行，但保留所有Markdown格式
        text_content = re.sub(r'\n{3,}', '\n\n', md_content)

        return {
            "content": text_content.strip(),
            "metadata": {
                "title": front_matter.get("title", title),
                "tags": front_matter.get("tags", []),
                "author": front_matter.get("author", "")
            },
            "source": os.path.basename(file_path),
            "format": "markdown"
        }

    def extract_text(self, content: Any) -> str:
        """提取纯文本"""
        return content


class TextParser(DocumentParser):
    """纯文本解析器"""

    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析文本文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        # 提取第一行作为标题
        lines = content.split("\n")
        title = lines[0][:100] if lines else ""

        return {
            "content": content,
            "metadata": {
                "title": title,
                "lines": len(lines)
            },
            "source": os.path.basename(file_path),
            "format": "txt"
        }

    def extract_text(self, content: Any) -> str:
        """提取纯文本"""
        return content


class DocumentParserFactory:
    """文档解析工厂"""

    PARSERS = {
        ".pdf": PDFParser,
        ".docx": DocxParser,
        ".doc": DocxParser,
        ".html": HTMLParser,
        ".htm": HTMLParser,
        ".md": MarkdownParser,
        ".txt": TextParser,
    }

    @classmethod
    def get_parser(cls, file_path: str) -> DocumentParser:
        """根据文件扩展名获取对应的解析器"""
        ext = Path(file_path).suffix.lower()
        parser_class = cls.PARSERS.get(ext)

        if parser_class is None:
            raise ValueError(f"不支持的文件格式: {ext}")

        return parser_class()

    @classmethod
    def parse(cls, file_path: str) -> Dict[str, Any]:
        """解析任意支持的文档"""
        parser = cls.get_parser(file_path)
        return parser.parse(file_path)


class TextCleaner:
    """文本清洗工具"""

    def __init__(self):
        self.url_pattern = re.compile(r'https?://\S+')
        self.email_pattern = re.compile(r'\S+@\S+\.\S+')
        self.phone_pattern = re.compile(r'\d{3}-?\d{4}-?\d{4}|\d{11}')
        # Markdown 特殊字符: # ` * _ [ ] ( ) - | > + =
        self.special_chars_pattern = re.compile(r'[^\w\s\u4e00-\u9fff.,!?;:，。！？；：""''（）《》【】\-\(\)\[\]\<\>#`\*_\[\]\(\)\-|\>+=]')
        self.multiple_spaces_pattern = re.compile(r'[ \t]+')  # 只匹配空格和制表符，不匹配换行
        self.multiple_newlines_pattern = re.compile(r'\n{3,}')

    def clean(self, text: str, remove_urls: bool = True, remove_emails: bool = True) -> str:
        """清洗文本"""
        cleaned = text

        if remove_urls:
            cleaned = self.url_pattern.sub('[链接]', cleaned)

        if remove_emails:
            cleaned = self.email_pattern.sub('[邮箱]', cleaned)

        # 移除特殊字符（保留中文、英文、数字和常用标点）
        cleaned = self.special_chars_pattern.sub(' ', cleaned)

        # 标准化空格
        cleaned = self.multiple_spaces_pattern.sub(' ', cleaned)

        # 标准化换行
        cleaned = self.multiple_newlines_pattern.sub('\n\n', cleaned)

        # 去除首尾空白
        cleaned = cleaned.strip()

        return cleaned

    def normalize_whitespace(self, text: str) -> str:
        """标准化空白字符"""
        return self.multiple_spaces_pattern.sub(' ', text)

    def remove_control_characters(self, text: str) -> str:
        """移除控制字符"""
        # 移除除\n\r\t外的控制字符
        return ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')

    def standardize_terminology(self, text: str, glossary: Dict[str, str] = None) -> str:
        """标准化术语"""
        if glossary is None:
            # 网络安全领域常用术语标准化
            glossary = {
                "SQL Injection": "SQL注入",
                "XSS": "跨站脚本攻击",
                "CSRF": "跨站请求伪造",
                "SSRF": "服务端请求伪造",
                "RCE": "远程代码执行",
                "webshell": "WebShell",
                "DDoS": "分布式拒绝服务攻击",
                "APT": "高级持续性威胁",
            }

        for eng, chn in glossary.items():
            text = re.sub(rf'\b{eng}\b', chn, text, flags=re.IGNORECASE)

        return text


# 全局实例
document_parser = DocumentParserFactory()
text_cleaner = TextCleaner()


def parse_document(file_path: str, clean_text: bool = True) -> Dict[str, Any]:
    """解析文档的便捷函数"""
    result = document_parser.parse(file_path)

    if clean_text and result.get("content"):
        result["content"] = text_cleaner.clean(result["content"])

    return result


def parse_documents_batch(file_paths: List[str], clean_text: bool = True) -> List[Dict[str, Any]]:
    """批量解析文档"""
    results = []
    for path in file_paths:
        try:
            result = parse_document(path, clean_text)
            results.append(result)
        except Exception as e:
            print(f"解析文件失败 {path}: {e}")
            results.append({
                "source": os.path.basename(path),
                "error": str(e),
                "content": "",
                "metadata": {}
            })
    return results