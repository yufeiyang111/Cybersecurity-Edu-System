# -*- coding: utf-8 -*-
"""把评测运行产物渲染成报告：自包含 HTML（双击即看）与 Word 文档（python-docx）。

输入：
- 汇总报告 ``backend/rag_report_<tag>_<dataset>.json``
- 逐题明细 ``backend/data/eval_run_<tag>_<dataset>.jsonl``

输出：
- ``backend/data/eval_report_<tag>_<dataset>.html``
- ``backend/data/eval_report_<tag>_<dataset>.docx``（--format docx/both）

用法（在 backend/ 下）::

    venv\\Scripts\\python.exe -m app.scripts.render_eval_report --tag m27full --dataset curated
    venv\\Scripts\\python.exe -m app.scripts.render_eval_report --tag m27full --format both
"""
from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

BACKEND_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_DATA_DIR = BACKEND_ROOT / "data"
CORPUS_VERSION = "knowledge_embeddings-v1"


def build_argument_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="渲染评测结果报告（HTML / Word）")
    parser.add_argument("--tag", required=True)
    parser.add_argument("--dataset", default="curated")
    parser.add_argument(
        "--format",
        choices=("html", "docx", "both"),
        default="html",
        help="输出格式（默认 html；docx 需要 python-docx，已内置）",
    )
    return parser


def _pct(value: Any) -> str:
    try:
        return f"{float(value) * 100:.1f}%"
    except (TypeError, ValueError):
        return "-"


def _bar(label: str, value: float, total: int, color: str) -> str:
    width = int(value * 1000) / 10 if total else 0
    return (
        f'<div class="bar-row"><span class="bar-label">{label}</span>'
        f'<div class="bar-track"><div class="bar-fill" style="width:{width}%;background:{color}"></div></div>'
        f"<span class=\"bar-value\">{value}</span></div>"
    )


def render(rows: List[Dict[str, Any]], tag: str, dataset: str) -> str:
    total = len(rows)
    r1 = [r["retrieval_metrics"].get("recall_at_1") for r in rows]
    r5 = [r["retrieval_metrics"].get("recall_at_5") for r in rows]
    mrr = [r["retrieval_metrics"].get("mrr") for r in rows]

    def mean(values):
        nums = [v for v in values if isinstance(v, (int, float))]
        return sum(nums) / len(nums) if nums else 0.0

    supported = sum(1 for r in rows if r.get("execution", {}).get("answer_status") == "supported")
    hit1 = sum(1 for v in r1 if v == 1.0)
    fail_stages: Dict[str, int] = {}
    for r in rows:
        stage = r.get("failure_stage") or "none"
        fail_stages[stage] = fail_stages.get(stage, 0) + 1
    by_diff: Dict[str, List[Dict[str, Any]]] = {}
    for r in rows:
        by_diff.setdefault(r["difficulty"], []).append(r)

    cards = [
        ("总题数", f"{total}", ""),
        ("Recall@1", _pct(mean(r1)), f"命中 {hit1} 题"),
        ("Recall@5", _pct(mean(r5)), ""),
        ("MRR", f"{mean(mrr):.3f}", ""),
        ("Supported 率", _pct(supported / total) if total else "-", f"{supported}/{total}"),
    ]
    cards_html = "".join(
        f'<div class="card"><div class="card-value">{v}</div><div class="card-label">{k}</div>'
        f'<div class="card-sub">{sub}</div></div>'
        for k, v, sub in cards
    )

    colors = ["#2563eb", "#16a34a", "#d97706", "#7c3aed", "#dc2626"]
    diff_bars = "".join(
        _bar(
            f"{diff}（{len(rs)}题）",
            round(sum(1 for r in rs if r["retrieval_metrics"].get("recall_at_1") == 1.0) / len(rs), 4),
            len(rs),
            colors[i % len(colors)],
        )
        for i, (diff, rs) in enumerate(sorted(by_diff.items()))
    )
    stage_bars = "".join(
        _bar(stage, count / total, total, colors[i % len(colors)])
        for i, (stage, count) in enumerate(
            sorted(fail_stages.items(), key=lambda kv: -kv[1])
        )
    )

    table_rows = []
    for r in rows:
        exec_info = r.get("execution") or {}
        ok = r["retrieval_metrics"].get("recall_at_1") == 1.0
        cls = "ok" if ok and not r.get("failure_stage") else ("warn" if ok else "bad")
        top5 = ", ".join(str(x) for x in exec_info.get("candidate_top5", []))
        table_rows.append(
            "<tr class='" + cls + "'>"
            f"<td>{r['case_key']}</td><td class='q'>{r['query']}</td>"
            f"<td>{r['difficulty']}</td><td>{','.join(r['expected_document_ids'])}</td>"
            f"<td class='mono'>{top5}</td>"
            f"<td>{exec_info.get('answer_status') or '-'}</td>"
            f"<td>{r.get('failure_stage') or '-'}</td>"
            f"<td>{exec_info.get('retrieval_ms') or '-'}</td></tr>"
        )

    html = f"""<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="utf-8">
<title>RAG 评测报告 · {tag} · {dataset}</title>
<style>
body{{font-family:'Microsoft YaHei',sans-serif;margin:24px;color:#1e293b;background:#f8fafc}}
h1{{font-size:20px}} .meta{{color:#64748b;font-size:12px;margin-bottom:18px}}
.cards{{display:flex;gap:14px;flex-wrap:wrap;margin-bottom:22px}}
.card{{background:#fff;border-radius:10px;padding:14px 20px;box-shadow:0 1px 3px rgba(0,0,0,.08);min-width:130px}}
.card-value{{font-size:26px;font-weight:700;color:#2563eb}}
.card-label{{font-size:13px;margin-top:2px}} .card-sub{{font-size:11px;color:#94a3b8}}
.panel{{background:#fff;border-radius:10px;padding:16px 20px;box-shadow:0 1px 3px rgba(0,0,0,.08);margin-bottom:20px}}
.panel h2{{font-size:15px;margin:0 0 12px}}
.bar-row{{display:flex;align-items:center;gap:10px;margin:6px 0}}
.bar-label{{width:170px;font-size:12px;text-align:right;color:#475569}}
.bar-track{{flex:1;height:16px;background:#f1f5f9;border-radius:999px;overflow:hidden}}
.bar-fill{{height:100%;border-radius:999px}}
.bar-value{{width:64px;font-size:12px;font-weight:600}}
input#f{{padding:6px 10px;border:1px solid #cbd5e1;border-radius:6px;width:260px;margin-bottom:10px}}
label.tgl{{font-size:13px;margin-left:14px;cursor:pointer}}
table{{border-collapse:collapse;width:100%;font-size:12px;background:#fff}}
th,td{{border-bottom:1px solid #e2e8f0;padding:6px 8px;text-align:left;vertical-align:top}}
th{{background:#f1f5f9;position:sticky;top:0}}
tr.ok td:first-child{{border-left:3px solid #16a34a}}
tr.warn td:first-child{{border-left:3px solid #d97706}}
tr.bad td:first-child{{border-left:3px solid #dc2626}}
td.q{{max-width:360px}} td.mono{{font-family:Consolas,monospace;color:#475569}}
tr.hide{{display:none}}
</style></head><body>
<h1>RAG 评测报告 · {tag} / {dataset}</h1>
<div class="meta">生成时间：{datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC ｜ pipeline=v2 ｜ corpus={CORPUS_VERSION}</div>
<div class="cards">{cards_html}</div>
<div class="panel"><h2>各难度 Recall@1</h2>{diff_bars}</div>
<div class="panel"><h2>失败阶段分布（citation 为保守标记：同引用被多个论断复用）</h2>{stage_bars}</div>
<div class="panel">
<h2>逐题明细（共 {total} 条）</h2>
<input id="f" placeholder="输入关键词过滤 query / case_key…">
<label class="tgl"><input type="checkbox" id="onlyBad"> 只看未全对</label>
<table><thead><tr><th>case</th><th>query</th><th>难度</th><th>期望文档</th><th>召回Top5</th><th>状态</th><th>失败阶段</th><th>检索ms</th></tr></thead>
<tbody>{''.join(table_rows)}</tbody></table></div>
<script>
document.getElementById('f').addEventListener('input',apply);
document.getElementById('onlyBad').addEventListener('change',apply);
function apply(){{
 var q=document.getElementById('f').value.toLowerCase();
 var bad=document.getElementById('onlyBad').checked;
 document.querySelectorAll('tbody tr').forEach(function(tr){{
  var text=tr.innerText.toLowerCase();
  var isBad=tr.classList.contains('warn')||tr.classList.contains('bad');
  tr.classList.toggle('hide',(q&&!text.includes(q))||(bad&&!isBad));
 }});
}}
</script></body></html>"""
    return html


def _mean(values: List[Any]) -> float:
    nums = [v for v in values if isinstance(v, (int, float))]
    return sum(nums) / len(nums) if nums else 0.0


def _collect_stats(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    r1 = [r["retrieval_metrics"].get("recall_at_1") for r in rows]
    supported = sum(
        1 for r in rows if r.get("execution", {}).get("answer_status") == "supported"
    )
    stages: Dict[str, int] = {}
    for r in rows:
        stage = r.get("failure_stage") or "none"
        stages[stage] = stages.get(stage, 0) + 1
    by_diff: Dict[str, List[Dict[str, Any]]] = {}
    for r in rows:
        by_diff.setdefault(r["difficulty"], []).append(r)
    return {
        "total": len(rows),
        "recall_at_1": _mean(r1),
        "recall_at_5": _mean([r["retrieval_metrics"].get("recall_at_5") for r in rows]),
        "mrr": _mean([r["retrieval_metrics"].get("mrr") for r in rows]),
        "supported_rate": supported / len(rows) if rows else 0.0,
        "supported_count": supported,
        "hit1_count": sum(1 for v in r1 if v == 1.0),
        "stages": stages,
        "by_difficulty": by_diff,
    }


def render_docx(rows: List[Dict[str, Any]], summary: Dict[str, Any], tag: str, dataset: str, out_path: Path) -> None:
    """生成 Word 评测报告：指标表 + 分难度表现 + 失败分布 + 未命样例。"""
    from docx import Document
    from docx.shared import Pt

    stats = _collect_stats(rows)
    document = Document()
    title = document.add_heading(f"RAG 评测报告 · {tag} / {dataset}", level=0)
    for run in title.runs:
        run.font.size = Pt(20)
    document.add_paragraph(
        f"生成时间：{datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC ｜ "
        f"pipeline=v2 ｜ corpus={CORPUS_VERSION} ｜ 模型=MiniMax-M2.7"
    )

    document.add_heading("一、核心指标", level=1)
    metrics = summary.get("metrics", {})
    retrieval = metrics.get("retrieval", {})
    runtime = metrics.get("runtime", {})
    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "指标"
    header[1].text = "数值"
    header[2].text = "说明"
    metric_rows = [
        ("总题数", str(stats["total"]), "curated 数据集全量"),
        ("Recall@1", _pct(stats["recall_at_1"]), f"期望文档排名第一 {stats['hit1_count']} 题"),
        ("Recall@5", _pct(stats["recall_at_5"]), "前五召回"),
        ("MRR", f"{stats['mrr']:.3f}", "平均倒数排名"),
        ("NDCG@10", f"{retrieval.get('ndcg_at_10', 0):.3f}", "排序质量"),
        ("Supported 率", _pct(stats["supported_rate"]), f"{stats['supported_count']}/{stats['total']} 给出有据回答"),
        ("证据覆盖率", _pct(metrics.get("evidence", {}).get("expected_evidence_coverage")), "引用证据覆盖期望内容"),
        ("检索延迟 P50/P95", f"{runtime.get('retrieval_p50_ms', '-')} / {runtime.get('retrieval_p95_ms', '-')} ms", "-"),
        ("发布阻断项", str(len(summary.get("release_blockers", []))), "无硬性阻断即通过"),
    ]
    for name, value, note in metric_rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = value
        cells[2].text = note

    document.add_heading("二、分难度 Recall@1", level=1)
    diff_table = document.add_table(rows=1, cols=3)
    diff_table.style = "Light Grid Accent 1"
    head = diff_table.rows[0].cells
    head[0].text = "难度"
    head[1].text = "题数"
    head[2].text = "Recall@1"
    for diff in ("easy", "medium", "hard"):
        rs = stats["by_difficulty"].get(diff, [])
        if not rs:
            continue
        hit = sum(1 for r in rs if r["retrieval_metrics"].get("recall_at_1") == 1.0)
        cells = diff_table.add_row().cells
        cells[0].text = diff
        cells[1].text = str(len(rs))
        cells[2].text = _pct(hit / len(rs))

    document.add_heading("三、失败阶段分布", level=1)
    stage_table = document.add_table(rows=1, cols=3)
    stage_table.style = "Light Grid Accent 1"
    shead = stage_table.rows[0].cells
    shead[0].text = "阶段"
    shead[1].text = "题数"
    shead[2].text = "说明"
    stage_notes = {
        "none": "全部通过",
        "citation": "保守标记：同一引用编号被多个论断复用",
        "answer_status": "答案状态与预期不符（重点复盘）",
        "candidate": "执行异常",
    }
    for stage, count in sorted(stats["stages"].items(), key=lambda kv: -kv[1]):
        cells = stage_table.add_row().cells
        cells[0].text = stage
        cells[1].text = str(count)
        cells[2].text = stage_notes.get(stage, "")

    document.add_heading("四、未命中样例（Top 20）", level=1)
    misses = [r for r in rows if r["retrieval_metrics"].get("recall_at_1") != 1.0][:20]
    if misses:
        miss_table = document.add_table(rows=1, cols=4)
        miss_table.style = "Light Grid Accent 1"
        mhead = miss_table.rows[0].cells
        mhead[0].text = "case"
        mhead[1].text = "query"
        mhead[2].text = "期望文档"
        mhead[3].text = "召回 Top5"
        for r in misses:
            exec_info = r.get("execution") or {}
            top5 = ", ".join(str(x) for x in exec_info.get("candidate_top5", []))
            cells = miss_table.add_row().cells
            cells[0].text = r["case_key"]
            cells[1].text = r["query"]
            cells[2].text = ",".join(r["expected_document_ids"])
            cells[3].text = top5
    else:
        document.add_paragraph("无未命中样例。")

    document.add_paragraph(
        "说明：citation 阶段标记为保守判定（重复引用计数），不代表引用无效；"
        "完整逐题数据见同名 .jsonl 明细文件。"
    )
    document.save(out_path)


def main() -> None:
    args = build_argument_parser().parse_args()
    formats = ["html", "docx"] if args.format == "both" else [args.format]
    summary_path = BACKEND_ROOT / f"rag_report_{args.tag}_{args.dataset}.json"
    details_path = DEFAULT_DATA_DIR / f"eval_run_{args.tag}_{args.dataset}.jsonl"

    rows = [
        json.loads(line)
        for line in details_path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    outputs: Dict[str, str] = {}
    if "html" in formats:
        html_path = DEFAULT_DATA_DIR / f"eval_report_{args.tag}_{args.dataset}.html"
        html_path.write_text(render(rows, args.tag, args.dataset), encoding="utf-8", newline="\n")
        outputs["html"] = str(html_path)
    if "docx" in formats:
        docx_path = DEFAULT_DATA_DIR / f"eval_report_{args.tag}_{args.dataset}.docx"
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        render_docx(rows, summary, args.tag, args.dataset, docx_path)
        outputs["docx"] = str(docx_path)
    print(
        json.dumps(
            {
                "outputs": outputs,
                "cases": len(rows),
                "summary_source": str(summary_path),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()

